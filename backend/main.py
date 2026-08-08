import os
import sys
import io

# Hard fix for Windows UnicodeEncodeError on cp1252 terminals
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')


import traceback
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from typing import Optional

from langchain_groq import ChatGroq
try:
    from langchain_ollama import ChatOllama
except ImportError:
    ChatOllama = None
from langchain_huggingface import HuggingFaceEmbeddings

from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
import pdfplumber
import fitz
import easyocr
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from supabase import create_client, Client

load_dotenv()

app = FastAPI(title="CA Exam Prep AI API")

@app.get("/")
def read_root():
    """Redirect to the API documentation when visiting the root URL."""
    return RedirectResponse(url="/docs")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup Uploads Directory for Previews
UPLOAD_DIR = "./uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/api/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

# Lazy OCR Reader initialization
_ocr_reader = None

def get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        print("Initializing EasyOCR Model on demand...")
        _ocr_reader = easyocr.Reader(['en'], gpu=False)
    return _ocr_reader

# Initialize Supabase
url = os.getenv("SUPABASE_URL")
key = os.getenv("SUPABASE_KEY")
if not url or not key:
    print("Warning: Supabase credentials missing!")
supabase: Client | None = create_client(url, key) if url and key else None

# Initialize RAG Components
DB_DIR = "../database/chroma_db" if os.path.exists("../database/chroma_db") else "./database/chroma_db" if os.path.exists("./database/chroma_db") else "./chroma_db"
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
vectorstore = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)
retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

AVAILABLE_MODELS = [
    {"id": "groq:llama-3.1-8b", "name": "Llama 3.1 8B (Cloud)", "provider": "Groq Cloud", "ram": "0 MB RAM", "badge": "⚡ Ultra Fast"},
    {"id": "groq:llama-3.3-70b", "name": "Llama 3.3 70B (Cloud)", "provider": "Groq Cloud", "ram": "0 MB RAM", "badge": "🧠 Ultra Smart"},
    {"id": "groq:mixtral-8x7b", "name": "Mixtral 8x7B (Cloud)", "provider": "Groq Cloud", "ram": "0 MB RAM", "badge": "🌟 High Logic"},
    {"id": "ollama:llama3.2:1b", "name": "Llama 3.2 1B (Local)", "provider": "Ollama Local", "ram": "1.3 GB RAM", "badge": "💻 Offline 1B"},
    {"id": "ollama:llama3.2:3b", "name": "Llama 3.2 3B (Local)", "provider": "Ollama Local", "ram": "2.2 GB RAM", "badge": "💻 Offline 3B"},
    {"id": "ollama:mistral:7b", "name": "Mistral 7B (Local)", "provider": "Ollama Local", "ram": "4.4 GB RAM", "badge": "⚡ 16GB RAM"},
    {"id": "ollama:qwen2.5:7b", "name": "Qwen 2.5 7B (Local)", "provider": "Ollama Local", "ram": "4.7 GB RAM", "badge": "🎯 16GB RAM"},
    {"id": "ollama:deepseek-r1:8b", "name": "DeepSeek R1 8B (Local)", "provider": "Ollama Local", "ram": "4.9 GB RAM", "badge": "🧠 16GB RAM"}
]

# Groq Cloud API Setup
groq_api_key = os.getenv("GROQ_API_KEY")

def get_llm_instance(model_id: Optional[str] = None):
    model_id = model_id or os.getenv("DEFAULT_MODEL", "groq:llama-3.1-8b")
    
    if model_id.startswith("ollama:"):
        ollama_name = model_id.replace("ollama:", "")
        ollama_base = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        if ChatOllama is not None:
            return ChatOllama(model=ollama_name, temperature=0.2, base_url=ollama_base)
        try:
            from langchain_ollama import ChatOllama as LCOllama
            return LCOllama(model=ollama_name, temperature=0.2, base_url=ollama_base)
        except Exception:
            from langchain_community.llms.ollama import Ollama
            return Ollama(model=ollama_name, temperature=0.2, base_url=ollama_base)

    
    # Groq Cloud fallback
    groq_model_name = "llama-3.1-8b-instant"
    if "70b" in model_id:
        groq_model_name = "llama-3.3-70b-versatile"
    elif "mixtral" in model_id:
        groq_model_name = "mixtral-8x7b-32768"
        
    return ChatGroq(
        model=groq_model_name,
        temperature=0.2,
        api_key=groq_api_key
    )

llm = get_llm_instance()

template = """You are a friendly, empathetic, and highly encouraging human tutor helping a student prepare for their CA exams (Foundation, Intermediate, Final).
You should communicate in a natural, conversational, and human-like tone. Use encouraging words, be polite, and sound like a real mentor who genuinely cares about the student's success. 
Even though you are friendly, your goal is still to provide highly accurate and clear answers based ONLY on the provided study materials.

Instructions:
1. Act human. Start your responses warmly (e.g., "Great question!", "I completely understand why that's confusing", "Let's look at this together"). 
2. The user may make spelling mistakes or typos in their question. Please infer their true intent and answer accordingly.
3. Use the provided context to answer the question. Break down complex accounting concepts into easy-to-understand human language.
4. If the answer is completely absent from the context, gently state that you don't have that specific document in the database, but provide general guidance based on official ICAI guidelines.
5. MANDATORY SOURCE CITATIONS: At the end of your answer, list the exact source document names and topics referenced under a "📚 Source References & Links" heading.

Context: {context}

Question: {question}

Answer:"""
prompt = ChatPromptTemplate.from_template(template)

def format_docs(docs):
    return "\n\n".join(f"[Source PDF: {doc.metadata.get('source', 'Unknown')}]\n{doc.page_content}" for doc in docs)

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    model_id: Optional[str] = None

@app.get("/api/models")
async def get_models():
    return {"models": AVAILABLE_MODELS}

@app.get("/api/sessions")
async def get_sessions():
    if not supabase: return {"sessions": []}
    try:
        response = supabase.table("sessions").select("*").order("created_at", desc=True).execute()
        return {"sessions": response.data or []}
    except Exception as e:
        print(f"Supabase fetch sessions notice (non-fatal): {e}")
        return {"sessions": []}

@app.get("/api/sessions/{session_id}")
async def get_session_messages(session_id: str):
    if not supabase: return {"messages": []}
    try:
        response = supabase.table("messages").select("*").eq("session_id", session_id).order("created_at", desc=False).execute()
        return {"messages": response.data or []}
    except Exception as e:
        print(f"Supabase fetch messages notice (non-fatal): {e}")
        return {"messages": []}

@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str):
    if not supabase: return {"message": "Session removed"}
    try:
        supabase.table("messages").delete().eq("session_id", session_id).execute()
        response = supabase.table("sessions").delete().eq("id", session_id).execute()
        return {"message": "Session deleted successfully"}
    except Exception as e:
        print(f"Supabase delete session notice (non-fatal): {e}")
        return {"message": "Session deleted"}


GREETINGS = {"hi", "hii", "hiii", "hello", "hey", "hlo", "greetings", "good morning", "good afternoon", "good evening"}

@app.post("/api/chat")
async def chat(request: ChatRequest):
    try:
        session_id = request.session_id
        
        # If no session, try creating one safely
        if not session_id and supabase:
            try:
                title = request.message[:30] + "..." if len(request.message) > 30 else request.message
                session_resp = supabase.table("sessions").insert({"title": title}).execute()
                if isinstance(session_resp.data, list) and len(session_resp.data) > 0:
                    first_session = session_resp.data[0]
                    if isinstance(first_session, dict):
                        session_id = first_session.get("id")
            except Exception as se:
                print(f"Supabase session creation error (non-fatal): {se}")
        
        # Save user message safely
        if session_id and supabase:
            try:
                supabase.table("messages").insert({
                    "session_id": session_id,
                    "role": "user",
                    "content": request.message
                }).execute()
            except Exception as me:
                print(f"Supabase user msg insert error (non-fatal): {me}")

        # Handle simple greetings concisely without RAG clutter
        clean_msg = request.message.strip().lower().rstrip("!.?")
        if clean_msg in GREETINGS or (len(clean_msg) <= 6 and clean_msg.startswith("hi")) or clean_msg in ["hello", "hey", "hlo"]:
            greeting_resp = "Hello! 😊 How can I help you with your CA exam preparation today? Feel free to ask any doubt about Tax, Costing, Auditing, or Law!"
            if session_id and supabase:
                try:
                    supabase.table("messages").insert({
                        "session_id": session_id,
                        "role": "ai",
                        "content": greeting_resp
                    }).execute()
                except Exception:
                    pass
            return {"response": greeting_resp, "session_id": session_id}

        # Retrieve relevant docs for exact source citations
        retrieved_docs = retriever.invoke(request.message)
        sources = list(set([doc.metadata.get('source') for doc in retrieved_docs if doc.metadata.get('source')]))
        context_str = format_docs(retrieved_docs)

        # Resolve target model dynamically per request
        target_model = request.model_id or os.getenv("DEFAULT_MODEL", "groq:llama-3.1-8b")
        model_fallback_notice = ""

        try:
            selected_llm = get_llm_instance(target_model)
            dynamic_rag_chain = prompt | selected_llm | StrOutputParser()
            response = dynamic_rag_chain.invoke({"context": context_str, "question": request.message})
        except Exception as model_err:
            print(f"Model invocation failed for '{target_model}', falling back to Groq Cloud LLM: {model_err}")
            fallback_llm = get_llm_instance("groq:llama-3.1-8b")
            fallback_chain = prompt | fallback_llm | StrOutputParser()
            response = fallback_chain.invoke({"context": context_str, "question": request.message})
            model_fallback_notice = f"\n\n*(Note: Local model '{target_model}' encountered an issue. Automatically switched to Groq Llama 3.1 8B Cloud AI).* "

        if model_fallback_notice:
            response += model_fallback_notice

        # Attach exact source document links if missing
        if sources and "Source References" not in response:
            source_links = "\n\n---\n### 📚 Source References & Links\n" + "\n".join([f"- {s}" for s in sources])
            response += source_links
        
        # Save AI response safely
        if session_id and supabase:
            try:
                supabase.table("messages").insert({
                    "session_id": session_id,
                    "role": "ai",
                    "content": response
                }).execute()
            except Exception as ae:
                print(f"Supabase AI msg insert error (non-fatal): {ae}")
            
        return {"response": response, "session_id": session_id}

    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Chat error")

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff",
    ".txt", ".md", ".csv", ".json",
    ".docx",
    ".webm", ".wav", ".mp3", ".m4a", ".ogg"
}

def extract_text_from_file(file_path: str, filename: str) -> tuple[str, bool]:
    """
    Extracts text from various file formats.
    Returns (extracted_text, needs_background_ocr).
    """
    ext = os.path.splitext(filename)[1].lower()
    
    # 1. Text & Markdown files
    if ext in [".txt", ".md", ".csv", ".json"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), False
        except Exception as e:
            print(f"Error reading text file {filename}: {e}")
            return "", False

    # 2. Word Documents (.docx)
    if ext == ".docx":
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text), False
        except Exception as e:
            print(f"Error reading docx file {filename}: {e}")
            return "", False

    # 3. Direct Images & Handwritten Notes / Scans (.png, .jpg, .jpeg, .webp, .bmp, .tiff)
    if ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
        try:
            print(f"Running EasyOCR on image {filename}...")
            results = get_ocr_reader().readtext(file_path, detail=0)
            text = "\n".join(results)
            return text, False
        except Exception as e:
            print(f"Error running OCR on image {filename}: {e}")
            return "", False

    # 4. Voice Audio Clips (.webm, .wav, .mp3, .m4a, .ogg)
    if ext in [".webm", ".wav", ".mp3", ".m4a", ".ogg"]:
        print(f"Voice audio file {filename} uploaded and saved.")
        return f"[Voice Recording Uploaded: {filename}]", False

    # 5. PDF Documents (.pdf)
    if ext == ".pdf":
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"Could not extract text via pdfplumber from {filename}: {e}")
            
        if text.strip():
            return text, False
        else:
            # Scanned / handwritten PDF -> trigger background OCR
            return "", True

    return "", False


def process_ocr_background(file_path: str, filename: str):
    print(f"Starting background OCR processing for {filename}...")
    try:
        ext = os.path.splitext(filename)[1].lower()
        text = ""
        
        if ext == ".pdf":
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")
                result = get_ocr_reader().readtext(io.BytesIO(img_bytes), detail=0)
                text += " ".join(result) + "\n"

            doc.close()
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            results = get_ocr_reader().readtext(file_path, detail=0)
            text = "\n".join(results)
            
        if text.strip():
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = text_splitter.split_text(text)
            docs = [Document(page_content=chunk, metadata={"source": filename}) for chunk in chunks]
            vectorstore.add_documents(docs)
            print(f"Background OCR finished for {filename}. Ingested {len(chunks)} chunks.")
        else:
            print(f"Background OCR finished for {filename}, but no text was found.")
    except Exception:
        print(f"Background OCR failed for {filename}:")
        traceback.print_exc()

@app.post("/api/upload")
async def upload_document(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Invalid request: File name is missing.")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format '{ext}'. Supported formats: PDF, Images (PNG, JPG, WEBP, etc.), TXT, MD, DOCX."
        )
    
    try:
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        content = await file.read()
        with open(file_path, "wb") as buffer:
            buffer.write(content)
            
        text, needs_bg_ocr = extract_text_from_file(file_path, file.filename)
        
        if needs_bg_ocr:
            print(f"No text found in {file.filename}, scheduling OCR in background...")
            background_tasks.add_task(process_ocr_background, file_path, file.filename)
            return {
                "message": f"Success! {file.filename} is a scanned document or image. It is being processed and learned via OCR in the background.",
                "url": f"/api/uploads/{file.filename}"
            }
            
        if not text.strip():
            raise HTTPException(status_code=400, detail=f"No text could be extracted from {file.filename}.")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        
        docs = [Document(page_content=chunk, metadata={"source": file.filename}) for chunk in chunks]
        vectorstore.add_documents(docs)
        
        return {
            "message": f"Success! Successfully extracted text and trained on {len(chunks)} chunks from {file.filename}.",
            "url": f"/api/uploads/{file.filename}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")
